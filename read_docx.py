#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script để đọc file Word (.docx) và trích xuất text + hình ảnh
"""

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
except ImportError:
    print("[ERROR] Can cai dat thu vien python-docx:")
    print("   pip install python-docx")
    sys.exit(1)

def extract_images_from_docx(docx_path, output_dir="extracted_images"):
    """Trích xuất hình ảnh từ file .docx"""
    import zipfile
    import shutil
    
    # Tạo thư mục output
    output_path = Path(output_dir)
    output_path.mkdir(exist_ok=True)
    
    # File .docx thực chất là file ZIP
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        # Tìm tất cả các file hình ảnh trong thư mục word/media/
        image_files = [f for f in zip_ref.namelist() if f.startswith('word/media/')]
        
        extracted_images = []
        for img_file in image_files:
            # Lấy tên file
            img_name = os.path.basename(img_file)
            # Lưu vào thư mục output
            output_file = output_path / img_name
            with zip_ref.open(img_file) as source, open(output_file, 'wb') as target:
                shutil.copyfileobj(source, target)
            extracted_images.append(str(output_file))
            print(f"  [OK] Da trich xuat: {img_name}")
    
    return extracted_images

def read_docx_content(docx_path):
    """Đọc nội dung từ file .docx"""
    doc = Document(docx_path)
    
    content = []
    
    # Đọc từng paragraph
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            content.append(text)
    
    # Đọc từng table
    for table in doc.tables:
        content.append("\n[TABLE]")
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            if any(row_data):  # Chỉ thêm nếu có dữ liệu
                content.append(" | ".join(row_data))
        content.append("[/TABLE]\n")
    
    return "\n".join(content)

def main():
    # Set UTF-8 encoding cho Windows
    if sys.platform == 'win32':
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')
    
    docx_path = Path("backend/public/NotePBL.docx")
    
    if not docx_path.exists():
        print(f"[ERROR] Khong tim thay file: {docx_path}")
        sys.exit(1)
    
    print(f"[INFO] Dang doc file: {docx_path}\n")
    
    # Trích xuất hình ảnh
    print("[INFO] Dang trich xuat hinh anh...")
    images = extract_images_from_docx(docx_path)
    print(f"[OK] Da trich xuat {len(images)} hinh anh vao thu muc 'extracted_images/'\n")
    
    # Đọc nội dung text
    print("[INFO] Dang doc noi dung van ban...\n")
    content = read_docx_content(docx_path)
    
    # Hiển thị nội dung
    print("=" * 80)
    print("NOI DUNG VAN BAN:")
    print("=" * 80)
    print(content)
    print("=" * 80)
    
    # Lưu vào file text
    output_text = Path("extracted_text.txt")
    with open(output_text, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"\n[OK] Da luu noi dung text vao: {output_text}")

if __name__ == "__main__":
    main()

